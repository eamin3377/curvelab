"""DOCX export: an editable Word report built with python-docx.

Mirrors the PDF structure with real Word styles (Heading 1/2, styled
tables, embedded chart image) so the document can be tuned before
submission.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.engines.export.base import ExportedFile, ExportPayload, exporter, fmt_num

INDIGO_DARK = RGBColor(0x31, 0x2E, 0x81)
DOC_TABLE_ROWS = 40


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Append a grid-styled table from string rows (first row = header)."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True


@exporter("docx")
def build_docx(payload: ExportPayload) -> ExportedFile:
    """Render the fit as a Word document and return its bytes."""
    fit = payload.fit
    meta = payload.meta
    doc = Document()

    title = doc.add_heading(meta.title, level=0)
    for run in title.runs:
        run.font.color.rgb = INDIGO_DARK

    subtitle = doc.add_paragraph("Least Squares Curve Fitting — Numerical Methods Project")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_lines = [
        f"Institution: {meta.institution}" if meta.institution else "",
        f"Course: {meta.course}",
        f"Author: {meta.author}" if meta.author else "",
        f"Student ID: {meta.student_id}" if meta.student_id else "",
        f"Date: {meta.date}",
        f"Model: {fit.model.value} (degree {fit.degree}) · Data points: {fit.n} · R²: {fmt_num(fit.metrics.r2)}",
    ]
    for line in info_lines:
        if line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Input Data", level=1)
    rows = [["i", "x", "y"]]
    for i, (xv, yv) in enumerate(zip(payload.x, payload.y), start=1):
        if i > DOC_TABLE_ROWS:
            rows.append(["…", f"first {DOC_TABLE_ROWS} of {fit.n} rows", ""])
            break
        rows.append([str(i), fmt_num(xv), fmt_num(yv)])
    _add_table(doc, rows)

    doc.add_heading("2. Summations", level=1)
    _add_table(doc, [["Quantity", "Value"]] + [[s.key, fmt_num(s.value)] for s in fit.summations])

    doc.add_heading("3. Normal Equations (A · c = b)", level=1)
    ne_rows = []
    for row, rhs in zip(fit.normal_equations.matrix, fit.normal_equations.vector):
        ne_rows.append([*(fmt_num(v) for v in row), "|", fmt_num(rhs)])
    _add_table(doc, ne_rows)

    doc.add_heading("4. Step-by-Step Solution", level=1)
    for step in fit.steps:
        doc.add_heading(f"Step {step.index} — {step.title}", level=2)
        doc.add_paragraph(step.description)

    if payload.chart_png:
        doc.add_heading("5. Regression Graph", level=1)
        doc.add_picture(io.BytesIO(payload.chart_png), width=Cm(16))

    doc.add_heading("6. Results", level=1)
    eq = doc.add_paragraph(fit.equation.plain)
    eq.runs[0].font.bold = True
    _add_table(
        doc,
        [["Coefficient", "Value"]]
        + [[c.name, fmt_num(c.value, 8)] for c in fit.coefficients],
    )

    doc.add_heading("7. Goodness of Fit", level=1)
    m = fit.metrics
    _add_table(
        doc,
        [["Metric", "Value"]]
        + [
            ["R²", fmt_num(m.r2, 8)],
            ["Adjusted R²", fmt_num(m.adj_r2, 8) if m.adj_r2 is not None else "—"],
            ["RMSE", fmt_num(m.rmse, 8)],
            ["MAE", fmt_num(m.mae, 8)],
            ["MSE", fmt_num(m.mse, 8)],
            ["SSE", fmt_num(m.sse, 8)],
            ["SST", fmt_num(m.sst, 8)],
        ],
    )

    buf = io.BytesIO()
    doc.save(buf)
    return ExportedFile(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    )
